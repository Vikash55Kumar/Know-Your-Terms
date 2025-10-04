import io
import re
import pdfplumber
import docx
import pytesseract
import fitz
from PIL import Image
from transformers import pipeline
from flask import Flask, request, jsonify, send_file
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import threading
from concurrent.futures import ThreadPoolExecutor
import time
import os
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# Flask app
app = Flask(__name__)

# BERT embedding model - much faster and more accurate
_embedding_model = None
_embedding_lock = threading.Lock()
_legal_clause_embeddings = None

def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        with _embedding_lock:
            if _embedding_model is None:  # Double-check locking
                print("Loading BERT embedding model...")
                _embedding_model = SentenceTransformer('all-MiniLM-L6-v2')  # Fast and accurate
    return _embedding_model

# Legal document clause templates for BERT similarity matching
LEGAL_CLAUSES = [
    "rental agreement lease terms tenant landlord property rent deposit",
    "loan agreement borrower lender repayment interest installment collateral security",
    "employment contract job offer internship terms salary compensation benefits",
    "service agreement vendor client payment terms deliverables milestone",
    "legal contract parties signature witness governing law jurisdiction",
    "termination clause notice period breach default consequences",
    "confidentiality agreement non-disclosure information proprietary data",
    "partnership agreement business terms profit sharing responsibilities",
    "licensing agreement intellectual property usage rights restrictions",
    "purchase agreement buyer seller goods services warranty liability"
]

def get_legal_embeddings():
    global _legal_clause_embeddings
    if _legal_clause_embeddings is None:
        model = get_embedding_model()
        _legal_clause_embeddings = model.encode(LEGAL_CLAUSES)
        print(f"Pre-computed embeddings for {len(LEGAL_CLAUSES)} legal clause templates")
    return _legal_clause_embeddings

# Enhanced keywords for heuristic scoring
SECTION_CUES = [
    "agreement", "security deposit", "rental period", "payment terms",
    "termination", "arbitration", "jurisdiction",
    "witness", "signatory", "governing law", "parties", "definitions",
    "probation period", "internship duration", "performance", 
    "salary", "compensation", "notice period", "work expectations",
    "attendance", "leaves", "certificate", "offer letter",
    "loan", "borrower", "lender", "repayment", "interest", "collateral", 
    "internship", "employment", "contract"
]

# Enhanced cache with better key generation
_heuristic_cache = {}
_classification_cache = {}

def safe_join_text(parts):
    return "\n".join([p for p in parts if p])

def smart_chunk_text(text, target_chunks=4):
    """Smart chunking that creates fewer, more meaningful chunks"""
    words = text.split()
    total_words = len(words)
    
    if total_words <= 400:
        return [text]  # Don't chunk small texts
    
    # Calculate optimal chunk size based on text length
    words_per_chunk = max(300, total_words // target_chunks)
    
    chunks = []
    for i in range(0, total_words, words_per_chunk):
        chunk = " ".join(words[i:i + words_per_chunk])
        if len(chunk.strip()) > 100:  # Only meaningful chunks
            chunks.append(chunk)
        if len(chunks) >= target_chunks:
            break
    
    return chunks

def enhanced_heuristic_score(text):
    """Enhanced heuristic with weighted scoring"""
    # Use better cache key
    text_sample = text[:2000] if len(text) > 2000 else text
    text_hash = hash(text_sample)
    
    if text_hash in _heuristic_cache:
        return _heuristic_cache[text_hash]
    
    if not text:
        return 0.0
    
    t = text.lower()
    
    # Weighted keyword scoring
    high_value_keywords = ["agreement", "contract", "parties", "signature", "terms"]
    medium_value_keywords = ["payment", "termination", "deposit", "duration", "governing"]
    standard_keywords = [k for k in SECTION_CUES if k not in high_value_keywords + medium_value_keywords]
    
    score = 0.0
    total_weight = 0.0
    
    # High value keywords (weight: 3)
    for keyword in high_value_keywords:
        if re.search(r'\b' + re.escape(keyword) + r'\b', t):
            score += 3.0
        total_weight += 3.0
    
    # Medium value keywords (weight: 2)
    for keyword in medium_value_keywords:
        if re.search(r'\b' + re.escape(keyword) + r'\b', t):
            score += 2.0
        total_weight += 2.0
    
    # Standard keywords (weight: 1)
    for keyword in standard_keywords:
        if re.search(r'\b' + re.escape(keyword) + r'\b', t):
            score += 1.0
        total_weight += 1.0
    
    final_score = score / max(1, total_weight)
    _heuristic_cache[text_hash] = final_score
    
    # Manage cache size
    if len(_heuristic_cache) > 50:
        # Remove oldest entries
        oldest_keys = list(_heuristic_cache.keys())[:25]
        for key in oldest_keys:
            del _heuristic_cache[key]
    
    return final_score

def classify_single_chunk_fast(chunk):
    """Ultra-fast BERT embedding classification with caching"""
    chunk_hash = hash(chunk[:500])  # Use first 500 chars for cache
    
    if chunk_hash in _classification_cache:
        return _classification_cache[chunk_hash]
    
    try:
        # Get embedding for the chunk
        model = get_embedding_model()
        chunk_embedding = model.encode([chunk])
        
        # Get legal clause embeddings
        legal_embeddings = get_legal_embeddings()
        
        # Calculate cosine similarity
        similarities = cosine_similarity(chunk_embedding, legal_embeddings)[0]
        max_similarity = float(np.max(similarities))
        
        # Cache the result
        _classification_cache[chunk_hash] = max_similarity
        
        # Manage cache size
        if len(_classification_cache) > 100:
            oldest_keys = list(_classification_cache.keys())[:50]
            for key in oldest_keys:
                del _classification_cache[key]
        
        return max_similarity
        
    except Exception as e:
        print(f"BERT Classification error: {e}")
        return 0.0

def classify_agreement_ultra_optimized(text):
    """Ultra optimized BERT embedding classification"""
    details = {
        "chunks": 0,
        "votes": 0,
        "vote_ratio": 0.0,
        "heuristic": 0.0,
        "avg_chunk_score": 0.0,
        "reason": "",
        "processing_time": 0.0
    }
    
    start_time = time.time()

    if not text.strip():
        details["reason"] = "empty_text"
        return False, details

    # Step 1: Quick rejection of obvious non-legal documents
    t = text.lower()
    non_legal_indicators = [
        "target audience", "hackathon", "pitch deck", "business model", 
        "use case", "pain point", "ai solution", "grammarly for", 
        "would you like me to", "let's start with", "perfect, now you're"
    ]
    
    if any(indicator in t for indicator in non_legal_indicators):
        details.update({
            "reason": "non_legal_document_detected",
            "processing_time": round(time.time() - start_time, 2)
        })
        return False, details
    
    # Step 2: Enhanced heuristic check
    heur = enhanced_heuristic_score(text)
    details["heuristic"] = round(heur, 3)
    
    # Step 2: Precise document type detection (avoid false positives)
    t = text.lower()
    
    # Strong document type indicators with context
    if ("loan agreement" in t and ("borrower" in t or "lender" in t)) or \
       (("this loan agreement" in t or "loan contract" in t) and "repayment" in t):
        details.update({
            "chunks": 1,
            "votes": 1,
            "vote_ratio": 1.0,
            "avg_chunk_score": 0.95,
            "reason": "loan_agreement_detected",
            "processing_time": round(time.time() - start_time, 2)
        })
        return True, details
    
    # More precise internship detection - must have formal structure
    if (("internship offer letter" in t or "internship agreement" in t) and 
        ("terms & conditions" in t or "stipend" in t or "duration" in t)) or \
       ("offer you an internship" in t and "company" in t):
        details.update({
            "chunks": 1,
            "votes": 1,
            "vote_ratio": 1.0,
            "avg_chunk_score": 0.95,
            "reason": "internship_agreement_detected",
            "processing_time": round(time.time() - start_time, 2)
        })
        return True, details
    
    # Rental agreement detection
    if ("rental agreement" in t or "lease agreement" in t) and \
       ("tenant" in t or "landlord" in t or "rent" in t):
        details.update({
            "chunks": 1,
            "votes": 1,
            "vote_ratio": 1.0,
            "avg_chunk_score": 0.95,
            "reason": "rental_agreement_detected",
            "processing_time": round(time.time() - start_time, 2)
        })
        return True, details
    
    # Employment contract detection
    if ("employment contract" in t or "job offer" in t) and \
       ("salary" in t or "position" in t or "employment" in t):
        details.update({
            "chunks": 1,
            "votes": 1,
            "vote_ratio": 1.0,
            "avg_chunk_score": 0.95,
            "reason": "employment_agreement_detected",
            "processing_time": round(time.time() - start_time, 2)
        })
        return True, details
    
    # Step 3: BERT embedding classification for ambiguous cases
    chunks = smart_chunk_text(text, target_chunks=2)  # Reduced to 2 chunks for speed
    details["chunks"] = len(chunks)
    
    if not chunks:
        details["reason"] = "no_chunks"
        return False, details

    # Step 4: Ultra-fast BERT similarity processing
    votes = 0
    per_chunk_scores = []
    BERT_THRESHOLD = 0.6  # BERT similarity threshold
    
    for i, chunk in enumerate(chunks):
        try:
            similarity_score = classify_single_chunk_fast(chunk)
            per_chunk_scores.append(similarity_score)
            
            if similarity_score >= BERT_THRESHOLD:
                votes += 1
            
            # Early stopping for clear positive/negative cases
            if similarity_score >= 0.8:  # Very high confidence
                details.update({
                    "votes": 1,
                    "vote_ratio": 1.0,
                    "avg_chunk_score": round(similarity_score, 3),
                    "reason": f"high_bert_confidence_{round(similarity_score, 2)}",
                    "processing_time": round(time.time() - start_time, 2)
                })
                return True, details
                
        except Exception as e:
            print(f"BERT classification error: {e}")
            per_chunk_scores.append(0.0)

    if not per_chunk_scores:
        # Fallback to heuristic
        if heur >= 0.3:
            details.update({
                "votes": 1,
                "vote_ratio": 1.0,
                "avg_chunk_score": heur,
                "reason": "bert_error_heuristic_fallback",
                "processing_time": round(time.time() - start_time, 2)
            })
            return True, details
        details["reason"] = "bert_error_low_heuristic"
        return False, details

    ratio = votes / len(per_chunk_scores)
    avg_score = sum(per_chunk_scores) / len(per_chunk_scores)
    
    details.update({
        "votes": votes,
        "vote_ratio": round(ratio, 3),
        "avg_chunk_score": round(avg_score, 3),
        "processing_time": round(time.time() - start_time, 2)
    })

    # Step 5: Combined BERT + heuristic decision
    bert_confidence = avg_score
    combined_score = (bert_confidence * 0.7) + (heur * 0.3)  # Higher weight to BERT
    
    # Accept if BERT confidence is high or combined score is good
    accept = bert_confidence >= 0.6 or combined_score >= 0.4 or (ratio >= 0.5 and heur >= 0.2)
    
    if not accept:
        details["reason"] = f"low_bert_score_{round(combined_score, 2)}"
    else:
        details["reason"] = f"accepted_bert_score_{round(combined_score, 2)}"
    
    return accept, details

# Keep extraction functions optimized
def extract_pdf_optimized(file_stream):
    """Optimized PDF extraction"""
    try:
        file_stream.seek(0)
        with pdfplumber.open(file_stream) as pdf:
            total_pages = len(pdf.pages)
            
            # More aggressive page limits for speed
            if total_pages <= 3:
                pages_to_process = total_pages
            elif total_pages <= 8:
                pages_to_process = min(5, total_pages)
            else:
                pages_to_process = min(8, total_pages)
            
            texts = []
            for i in range(pages_to_process):
                page_text = pdf.pages[i].extract_text()
                if page_text and len(page_text.strip()) > 50:
                    texts.append(page_text)
                    
                # Early return if we have sufficient content
                if len("\n".join(texts)) > 5000:
                    break
                    
            return safe_join_text(texts)
            
    except Exception as e:
        print(f"PDF extract error: {e}")
        return ""

def extract_docx_optimized(file_stream):
    """Optimized DOCX extraction with early stopping"""
    try:
        file_stream.seek(0)
        doc = docx.Document(io.BytesIO(file_stream.read()))
        
        texts = []
        total_chars = 0
        
        for paragraph in doc.paragraphs:
            if paragraph.text and len(paragraph.text.strip()) > 10:
                texts.append(paragraph.text)
                total_chars += len(paragraph.text)
                
                # Early stopping for large documents
                if total_chars > 20000:  # Limit to ~20KB
                    break
                    
        return "\n".join(texts)
    except Exception as e:
        print(f"DOCX extract error: {e}")
        return ""

def extract_image_optimized(file_stream):
    """Optimized image extraction"""
    try:
        file_stream.seek(0)
        img = Image.open(file_stream).convert("RGB")
        
        # More aggressive resizing for speed
        max_dimension = 2500
        if img.width > max_dimension or img.height > max_dimension:
            img.thumbnail((max_dimension, max_dimension), Image.Resampling.LANCZOS)
        elif img.width < 600:  # Upscale very small images
            scale_factor = 600 / img.width
            new_size = (int(img.width * scale_factor), int(img.height * scale_factor))
            img = img.resize(new_size, Image.Resampling.LANCZOS)
            
        # Faster OCR config
        return pytesseract.image_to_string(img, config='--psm 6 --oem 3')
    except Exception as e:
        print(f"Image extract error: {e}")
        return ""

# Routes remain the same but use ultra-optimized function
@app.route("/active", methods=["GET"])
def active():
    return "active"

@app.route("/uploads", methods=["POST"])
def api_upload():
    start_time = time.time()
    
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    filename = file.filename.lower()
    text = ""
    extraction_time = 0

    # Extract text with timing
    extract_start = time.time()
    if filename.endswith(".pdf"):
        text = extract_pdf_optimized(file.stream)
    elif filename.endswith(".docx"):
        text = extract_docx_optimized(file.stream)
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        text = extract_image_optimized(file.stream)
    else:
        return jsonify({"error": "Unsupported file type"}), 400
    
    extraction_time = round(time.time() - extract_start, 2)

    # Classify with ultra-optimized function
    classify_start = time.time()
    is_ok, details = classify_agreement_ultra_optimized(text)
    classify_time = round(time.time() - classify_start, 2)
    
    total_time = round(time.time() - start_time, 2)

    if not is_ok:
        return jsonify({
            "error": "Rejected: Not a valid agreement.",
            "details": {
                **details,
                "extraction_time": extraction_time,
                "classification_time": classify_time,
                "total_time": total_time
            }
        }), 400

    return jsonify({
        "filename": file.filename,
        "extracted_text": text,
        # "processing_details": {
        #     **details,
        #     "extraction_time": extraction_time,
        #     "classification_time": classify_time,
        #     "total_time": total_time,
        #     "text_length": len(text)
        # }
    })

@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    text = request.form.get("text", "")
    output = io.BytesIO()
    doc = SimpleDocTemplate(output)
    styles = getSampleStyleSheet()
    story = [Paragraph(line, styles["Normal"]) for line in text.split("\n")]
    doc.build(story)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="output.pdf")

@app.route("/export/docx", methods=["POST"])
def export_docx():
    text = request.form.get("text", "")
    output = io.BytesIO()
    d = docx.Document()
    for line in text.split("\n"):
        d.add_paragraph(line)
    d.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="output.docx")

if __name__ == "__main__":
    # Pre-load the BERT model and embeddings at startup
    print("Loading BERT embedding model...")
    start_load = time.time()
    get_embedding_model()
    get_legal_embeddings()
    load_time = round(time.time() - start_load, 2)
    print(f"BERT model and embeddings loaded successfully in {load_time}s!")
    
    port = int(os.environ.get("PORT", 8000))
    app.run(host="0.0.0.0", port=port, threaded=True)